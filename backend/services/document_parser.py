"""Document parser for multiple file formats.

Extracts structured text with heading detection for:
- PDF (PyMuPDF with font-size heading detection)
- DOCX (python-docx with style introspection)
- Markdown (native heading parsing)
- TXT (implicit heading detection)
"""

import re
import logging
from pathlib import Path

import fitz  # PyMuPDF
import docx
import markdown
from bs4 import BeautifulSoup

from models.document import DocumentSection, ParsedDocument

logger = logging.getLogger(__name__)


class DocumentParser:
    """Parses documents into structured sections with heading hierarchy."""

    SUPPORTED_TYPES = {".pdf", ".docx", ".md", ".txt"}

    def parse(self, file_path: str) -> ParsedDocument:
        """Parse a document file into structured sections.

        Args:
            file_path: Path to the document file.

        Returns:
            ParsedDocument with extracted sections.

        Raises:
            ValueError: If file type is unsupported.
            FileNotFoundError: If file doesn't exist.
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_TYPES:
            raise ValueError(f"Unsupported file type: {ext}")

        logger.info(f"Parsing {path.name} (type: {ext})")

        try:
            if ext == ".pdf":
                sections = self._parse_pdf(path)
            elif ext == ".docx":
                sections = self._parse_docx(path)
            elif ext == ".md":
                sections = self._parse_markdown(path)
            elif ext == ".txt":
                sections = self._parse_txt(path)
            else:
                sections = []

            raw_text = "\n\n".join(s.content for s in sections)

            return ParsedDocument(
                file_name=path.name,
                file_path=str(path),
                document_type=ext.lstrip("."),
                sections=sections,
                raw_text=raw_text,
                metadata={"file_size": path.stat().st_size},
            )

        except Exception as e:
            logger.error(f"Error parsing {path.name}: {e}")
            raise

    def _parse_pdf(self, path: Path) -> list[DocumentSection]:
        """Parse PDF using PyMuPDF with font-size-based heading detection."""
        sections: list[DocumentSection] = []
        doc = fitz.open(str(path))

        current_section_title = "Document Start"
        current_section_level = 1
        current_content_parts: list[str] = []

        for page_num in range(len(doc)):
            page = doc[page_num]
            blocks = page.get_text("dict", sort=True)["blocks"]

            for block in blocks:
                if "lines" not in block:
                    continue

                for line in block["lines"]:
                    line_text = ""
                    max_font_size = 0

                    for span in line["spans"]:
                        line_text += span["text"]
                        max_font_size = max(max_font_size, span["size"])

                    line_text = line_text.strip()
                    if not line_text:
                        continue

                    # Detect headings based on font size
                    heading_level = self._detect_pdf_heading_level(
                        max_font_size, line_text
                    )

                    if heading_level > 0:
                        # Save previous section
                        if current_content_parts:
                            content = "\n".join(current_content_parts).strip()
                            if content:
                                sections.append(
                                    DocumentSection(
                                        title=current_section_title,
                                        content=content,
                                        level=current_section_level,
                                    )
                                )

                        current_section_title = line_text
                        current_section_level = heading_level
                        current_content_parts = []
                    else:
                        current_content_parts.append(line_text)

        # Save last section
        if current_content_parts:
            content = "\n".join(current_content_parts).strip()
            if content:
                sections.append(
                    DocumentSection(
                        title=current_section_title,
                        content=content,
                        level=current_section_level,
                    )
                )

        doc.close()

        # If no headings were detected, treat entire doc as one section
        if not sections:
            doc2 = fitz.open(str(path))
            full_text = ""
            for page in doc2:
                full_text += page.get_text() + "\n"
            doc2.close()
            if full_text.strip():
                sections.append(
                    DocumentSection(
                        title="Document",
                        content=full_text.strip(),
                        level=1,
                    )
                )

        return sections

    def _detect_pdf_heading_level(self, font_size: float, text: str) -> int:
        """Detect heading level from font size.

        Returns 0 if not a heading, 1-3 for heading levels.
        """
        # Skip very short text or text that looks like page numbers
        if len(text) < 2 or text.isdigit():
            return 0

        # Skip very long lines (likely body text)
        if len(text) > 200:
            return 0

        if font_size >= 18:
            return 1
        elif font_size >= 14:
            return 2
        elif font_size >= 12.5:
            return 3

        return 0

    def _parse_docx(self, path: Path) -> list[DocumentSection]:
        """Parse DOCX using python-docx with style introspection."""
        doc = docx.Document(str(path))
        sections: list[DocumentSection] = []

        current_section_title = "Document Start"
        current_section_level = 1
        current_content_parts: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else ""

            heading_level = 0
            if style_name.startswith("Heading"):
                try:
                    heading_level = int(style_name.split()[-1])
                except (ValueError, IndexError):
                    heading_level = 1

            if heading_level > 0:
                # Save previous section
                if current_content_parts:
                    content = "\n".join(current_content_parts).strip()
                    if content:
                        sections.append(
                            DocumentSection(
                                title=current_section_title,
                                content=content,
                                level=current_section_level,
                            )
                        )

                current_section_title = text
                current_section_level = heading_level
                current_content_parts = []
            else:
                current_content_parts.append(text)

        # Save last section
        if current_content_parts:
            content = "\n".join(current_content_parts).strip()
            if content:
                sections.append(
                    DocumentSection(
                        title=current_section_title,
                        content=content,
                        level=current_section_level,
                    )
                )

        return sections

    def _parse_markdown(self, path: Path) -> list[DocumentSection]:
        """Parse Markdown with native heading detection."""
        text = path.read_text(encoding="utf-8", errors="replace")
        return self._split_by_markdown_headings(text)

    def _split_by_markdown_headings(self, text: str) -> list[DocumentSection]:
        """Split markdown text by headings into sections."""
        sections: list[DocumentSection] = []
        lines = text.split("\n")

        current_title = "Document Start"
        current_level = 1
        current_parts: list[str] = []

        heading_pattern = re.compile(r"^(#{1,6})\s+(.+)$")

        for i, line in enumerate(lines):
            match = heading_pattern.match(line)

            # Also check for setext-style headings (underline with === or ---)
            setext_level = 0
            if i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                if next_line and all(c == "=" for c in next_line):
                    setext_level = 1
                elif next_line and all(c == "-" for c in next_line) and len(next_line) >= 3:
                    setext_level = 2

            if match:
                # Save previous section
                if current_parts:
                    content = "\n".join(current_parts).strip()
                    if content:
                        sections.append(
                            DocumentSection(
                                title=current_title,
                                content=content,
                                level=current_level,
                            )
                        )
                hashes, title = match.groups()
                current_title = title.strip()
                current_level = len(hashes)
                current_parts = []
            elif setext_level > 0 and line.strip():
                # Save previous section
                if current_parts:
                    content = "\n".join(current_parts).strip()
                    if content:
                        sections.append(
                            DocumentSection(
                                title=current_title,
                                content=content,
                                level=current_level,
                            )
                        )
                current_title = line.strip()
                current_level = setext_level
                current_parts = []
            else:
                # Skip setext underlines
                prev_line = lines[i - 1].strip() if i > 0 else ""
                is_setext_underline = (
                    line.strip()
                    and (all(c == "=" for c in line.strip()) or
                         (all(c == "-" for c in line.strip()) and len(line.strip()) >= 3))
                    and prev_line
                )
                if not is_setext_underline:
                    current_parts.append(line)

        # Save last section
        if current_parts:
            content = "\n".join(current_parts).strip()
            if content:
                sections.append(
                    DocumentSection(
                        title=current_title,
                        content=content,
                        level=current_level,
                    )
                )

        return sections

    def _parse_txt(self, path: Path) -> list[DocumentSection]:
        """Parse TXT with implicit heading detection."""
        text = path.read_text(encoding="utf-8", errors="replace")
        sections: list[DocumentSection] = []

        # Try to detect markdown-style headings first
        if re.search(r"^#{1,6}\s+", text, re.MULTILINE):
            return self._split_by_markdown_headings(text)

        lines = text.split("\n")
        current_title = "Document"
        current_parts: list[str] = []

        for i, line in enumerate(lines):
            stripped = line.strip()

            # Detect ALL CAPS headings (at least 3 chars, not a common abbreviation)
            is_caps_heading = (
                stripped
                and stripped.isupper()
                and len(stripped) >= 3
                and len(stripped) <= 100
                and not stripped.isdigit()
                and " " in stripped or len(stripped) >= 5
            )

            if is_caps_heading:
                if current_parts:
                    content = "\n".join(current_parts).strip()
                    if content:
                        sections.append(
                            DocumentSection(
                                title=current_title,
                                content=content,
                                level=1,
                            )
                        )
                current_title = stripped.title()
                current_parts = []
            else:
                current_parts.append(line)

        # Save last section
        if current_parts:
            content = "\n".join(current_parts).strip()
            if content:
                sections.append(
                    DocumentSection(
                        title=current_title,
                        content=content,
                        level=1,
                    )
                )

        return sections
